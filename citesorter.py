import sys
import xml.etree.ElementTree as ET
import uuid
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Portions gratefully appropriated from https://github.com/JStrydhorst/pubmed-to-word 
# TODO - Pare down cloned code to just necessary parsing

topics = {
  'OUR DRUG HERE':0,
  'COMPETITORS – PSORIASIS':1,
  'PHOTO/TOPICAL PSORIASIS THERAPY':2,
  'PSORIASIS AND COMORBIDITIES':3,
  'GENETICS – PSORIASIS':4,
  'RWE, PRO, SUBPOPULATIONS, PHARMACOECONOMICS':5,
  'INTERLEUKINS, TARGETS, PATHOPHYSIOLOGY':6,
  'GENERAL – PSORIASIS':7,
  'GENERAL SAFETY':8,
  'COMPETITORS – PSORIATIC ARTHRITIS':9,
  'GENERAL – PSORIATIC ARTHRITIS':10,
  'SPONDYLOARTHRITIS':11
}

def parse_nbib(lines):
  # reads the lines of text until it finds a tag in the first four columns, then
  # concatenates the values until the start of the next tag
  nlines = 0
  tag = ''
  for line in lines:
    if len(tag) == 0:
      tag = line[0:4].strip()
      val = line[5:].strip()
      nlines += 1
    elif not line[0:4].strip():
      val = val + ' ' + line[5:].strip()
      nlines += 1
    else:
      break    
  return tag, val, nlines

def import_sources(filename):
  # converts the citation file to MSWord compatible bibliography xml file
  sources = ET.Element('b:Sources',
                       {"SelectedStyle" : "",
                        "xmlns:b" : "http://schemas.openxmlformats.org/officeDocument/2006/bibliography",
                        "xmlns" : "http://schemas.openxmlformats.org/officeDocument/2006/bibliography"})
  f = open(filename, 'rU')
  lines = f.readlines()
  line = 0
  while line < len(lines):
    tag, val, linesread = parse_nbib(lines[line:])
    line += linesread
    if tag == 'PMID':
      source = ET.SubElement(sources,'b:Source')
      sourcetag = ET.SubElement(source, 'b:Tag')
      sourcetype = ET.SubElement(source, 'b:SourceType')
      sourcetype.text = 'JournalArticle'
      sourceguid = ET.SubElement(source, 'b:Guid')
      sourceguid.text = '{' + str(uuid.uuid4()) + '}'
      sourcetitle = ET.SubElement(source, 'b:Title')
      sourceyear = ET.SubElement(source, 'b:Year')
      sourcejournal = ET.SubElement(source, 'b:JournalName')
      sourcepages = ET.SubElement(source, 'b:Pages')
      sourceauthor = ET.SubElement(source, 'b:Author')
      sourceauthor2 = ET.SubElement(sourceauthor, 'b:Author')
      sourcenames = ET.SubElement(sourceauthor2, 'b:NameList')
      sourcevolume = ET.SubElement(source, 'b:Volume')
    elif tag == 'TI':
      sourcetitle.text = val
    elif tag == 'TA':
      sourcejournal.text = val
    elif tag == 'DP':
      sourceyear.text = val[0:4]
    elif tag == 'FAU':
      match = re.search(r'^([\w\'-]+),\s*([\w\'-]+)\s*(.+)?' ,val)
      if match:
        person = ET.SubElement(sourcenames, 'b:Person')
        lastname = ET.SubElement(person, 'b:Last')
        lastname.text = match.group(1)
        if match.group(3):
          middlename = ET.SubElement(person, 'b:Middle')
          middlename.text = match.group(3)
        firstname = ET.SubElement(person, 'b:First')
        firstname.text = match.group(2)
        if not sourcetag.text:
          sourcetag.text = match.group(1)[0:3] + sourceyear.text[2:4]
    elif tag == 'VI':
      sourcevolume.text = val
    elif tag == 'PG':
      sourcepages.text = val
  f.close()
  tree = ET.ElementTree(sources)
  tree.write('sources.xml',encoding="utf-8",xml_declaration=True)

# EXTRACT DATA
#import_sources("citations.nbib")

# FORMAT DATA
# sorting based on keywords will go here probably

# EXPORT DATA
document = Document()
document.add_heading('Current Awareness and Competitor Update') #TODO - specify font and justifications
document.add_heading('DATE HERE 2017') #TODO - Uh, we can probably automate this
for title in topics:
  document.add_heading(title, level=1)
  # for citation in citations: #Where value = value of topic, post-sort. Need to assign these variables during parsing stage
  #   paragraph = document.add_paragraph()
  #   paragraph.add_run(authors)
  #   paragraph.add_run(title).bold = True
  #   paragraph.add_run(publication).italic = True
  #   paragraph.add_run(link)
  #   paragraph.add_run(abstract)

document.save('output.docx') #TODO - Automate this to have the current date as the filename




